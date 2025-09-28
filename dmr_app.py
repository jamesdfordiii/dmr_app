# dmr_app.py (TOP OF FILE)
import os
import re
from datetime import datetime, time

import numpy as np
import pandas as pd
import streamlit as st
import yfinance as yf
from zoneinfo import ZoneInfo

# ----------------- Page / Theme -----------------
st.set_page_config(page_title="Daily Market Report", layout="wide")

# ---- Simple password gate (place right after page_config) ----
def _password_gate():
    # Prefer Streamlit Secrets in the cloud; allow ENV var locally for convenience
    app_pw = st.secrets.get("app_password", None) or os.getenv("APP_PASSWORD")

    # If no password is configured, block by default (safer for cloud).
    # For local development, you can set APP_PASSWORD env var to skip secrets.
    if not app_pw:
        st.error("App password is not set. Add 'app_password' in Streamlit Secrets "
                 "or set APP_PASSWORD as an environment variable for local dev.")
        st.stop()

    # If already authenticated, let the script continue
    if st.session_state.get("pw_ok", False):
        return

    st.title("🔒 Daily Market Report")
    st.caption("This app is password-protected.")

    with st.form("pw_form", clear_on_submit=True):
        pw = st.text_input("Enter password", type="password")
        submitted = st.form_submit_button("Enter")

    if submitted:
        if pw == app_pw:
            st.session_state["pw_ok"] = True
            st.rerun()
        else:
            st.error("Incorrect password. Try again.")
            st.stop()
    else:
        st.stop()

_password_gate()
# ---- end password gate ----


# ----------------- Page / Theme -----------------
st.set_page_config(page_title="Daily Market Report", layout="wide")
st.title("📊 Daily Market Report")
st.caption("One-stop dashboard to create your DMR each morning.")

# ----------------- Helpers -----------------
NY = ZoneInfo("America/New_York")

def last_float(x) -> float:
    """Coerce Series/DataFrame/scalar to a single float (last non-NaN)."""
    try:
        if isinstance(x, pd.DataFrame):
            x = x.select_dtypes(include=[np.number]).iloc[-1].squeeze()
        if isinstance(x, (pd.Series, pd.Index)):
            x = pd.Series(x).dropna().iloc[-1]
        return float(x)
    except Exception:
        return float("nan")

@st.cache_data(show_spinner=False, ttl=120)
def fetch_daily_snapshot(symbols: list[str]) -> pd.DataFrame:
    """5d/1d snapshot for % change vs prior close."""
    out = {}
    for sym in symbols:
        try:
            df = yf.download(sym, period="5d", interval="1d", progress=False, auto_adjust=True)
            if df.empty:
                continue
            last = float(df["Close"].iloc[-1])
            prev = float(df["Close"].iloc[-2]) if len(df) > 1 else last
            chg = (last / prev - 1.0) * 100.0
            out[sym] = {"Last": round(last, 2), "Chg%": round(chg, 2)}
        except Exception as e:
            out[sym] = {"Last": np.nan, "Chg%": np.nan}
            st.warning(f"Daily snapshot failed for {sym}: {e}")
    return pd.DataFrame(out).T

@st.cache_data(show_spinner=False, ttl=60)
def fetch_premarket_stats(tickers: list[str]) -> pd.DataFrame:
    """
    Pre-market last, % change vs prev close, and pre-market volume.
    Uses 1m intraday with prepost=True and filters 04:00–09:30 ET for stocks/ETFs.
    For 24h assets (crypto/futures), uses since midnight ET.
    """
    today = datetime.now(tz=NY).date()
    pm_start = datetime.combine(today, time(4, 0), tzinfo=NY)
    pm_end   = datetime.combine(today, time(9, 30), tzinfo=NY)
    since_midnight = datetime.combine(today, time(0, 0), tzinfo=NY)

    rows = []
    for t in tickers:
        try:
            hist = yf.download(t, period="2d", interval="1m", prepost=True, progress=False, auto_adjust=True)
            if hist.empty:
                rows.append({"Symbol": t, "PM_Last": np.nan, "PM_%Chg": np.nan, "PM_Vol": np.nan})
                continue

            # make tz-aware in NY (yfinance usually returns UTC)
            if hist.index.tz is None:
                hist.index = hist.index.tz_localize("UTC").tz_convert(NY)
            else:
                hist.index = hist.index.tz_convert(NY)

            # previous close from last regular session
            d1d = yf.download(t, period="5d", interval="1d", progress=False, auto_adjust=True)
            prev_close = float(d1d["Close"].iloc[-2]) if len(d1d) > 1 else float(d1d["Close"].iloc[-1])

            # Market type heuristic
            is_24h = any(sfx in t.upper() for sfx in ["-USD", "=F"])

            if is_24h:
                window = hist.loc[hist.index >= since_midnight]
            else:
                window = hist.loc[(hist.index >= pm_start) & (hist.index < pm_end)]

            if window.empty:
                rows.append({"Symbol": t, "PM_Last": np.nan, "PM_%Chg": np.nan, "PM_Vol": 0})
                continue

            pm_last = float(window["Close"].iloc[-1])
            pm_vol  = int(window.get("Volume", pd.Series(dtype=float)).fillna(0).sum()) if "Volume" in window.columns else 0
            pm_chg  = ((pm_last / prev_close) - 1.0) * 100.0 if prev_close else np.nan

            rows.append({
                "Symbol": t,
                "PM_Last": round(pm_last, 2),
                "PM_%Chg": round(pm_chg, 2),
                "PM_Vol": pm_vol
            })
        except Exception as e:
            st.warning(f"Premarket fetch failed for {t}: {e}")
            rows.append({"Symbol": t, "PM_Last": np.nan, "PM_%Chg": np.nan, "PM_Vol": np.nan})

    return pd.DataFrame(rows).set_index("Symbol")

def combine_snapshot_and_pm(daily: pd.DataFrame, pm: pd.DataFrame) -> pd.DataFrame:
    df = daily.copy()
    for col in ["PM_Last", "PM_%Chg", "PM_Vol"]:
        df[col] = pm[col] if col in pm.columns else np.nan
    # helpful deltas
    try:
        df["Gap_vs_Close_%"] = (df["PM_Last"] / df["Last"] - 1.0) * 100.0
    except Exception:
        df["Gap_vs_Close_%"] = np.nan
    # nicer ordering
    cols = ["Last", "Chg%", "PM_Last", "PM_%Chg", "Gap_vs_Close_%", "PM_Vol"]
    df = df[[c for c in cols if c in df.columns]]
    # ensure PM_Vol numeric for formatting
    if "PM_Vol" in df.columns:
        df["PM_Vol"] = pd.to_numeric(df["PM_Vol"], errors="coerce")
    return df

def bias_from_symbol(symvals: dict) -> tuple[str, float, str]:
    """
    Determine "Risk-On/Risk-Off" bias using Dow futures (^YM=F) if present; else ^DJI.
    Returns (Bias, pct_change, used_symbol).
    """
    used = None
    pct = np.nan
    if "^YM=F" in symvals and not pd.isna(symvals["^YM=F"].get("Chg%", np.nan)):
        used = "^YM=F"; pct = float(symvals["^YM=F"]["Chg%"])
    elif "^DJI" in symvals and not pd.isna(symvals["^DJI"].get("Chg%", np.nan)):
        used = "^DJI"; pct = float(symvals["^DJI"]["Chg%"])
    bias = "Risk-On" if (not pd.isna(pct) and pct > 0) else "Risk-Off"
    return bias, pct, used or "—"

def parse_updown_text(pasted: str) -> pd.DataFrame:
    """
    Forgiving parser for lines like:
      AAPL  Upgrade  Neutral -> Buy  Firm: JPM  Analyst: Doe  PT: 210
      MSFT  Downgrade  Overweight->Equal-Weight  Morgan Stanley  PT: 400
    """
    rows = []
    for raw in pasted.splitlines():
        line = raw.strip()
        if not line:
            continue
        m = re.match(r"^([A-Z\.\-]+)\s+(.*)$", line)
        if not m:
            continue
        ticker, rest = m.group(1), m.group(2)

        action = "Upgrade" if re.search(r"upgrade", rest, re.I) else ("Downgrade" if re.search(r"downgrade", rest, re.I) else "")
        m2 = re.search(r"([A-Za-z\-\s]+)\s*[-–>]\s*([A-Za-z\-\s]+)", rest)
        old_new = (m2.group(1).strip(), m2.group(2).strip()) if m2 else ("", "")
        m3 = re.search(r"(Firm|Broker|House)\s*:\s*([A-Za-z0-9&\.\-\s]+)", rest, re.I)
        firm = m3.group(2).strip() if m3 else ""
        m4 = re.search(r"(Analyst)\s*:\s*([A-Za-z\.\-\s]+)", rest, re.I)
        analyst = m4.group(2).strip() if m4 else ""
        m5 = re.search(r"PT\s*:\s*([0-9]+(\.[0-9]+)?)", rest, re.I)
        pt = float(m5.group(1)) if m5 else np.nan

        rows.append({
            "Ticker": ticker,
            "Action": action,
            "OldRating": old_new[0],
            "NewRating": old_new[1],
            "Firm": firm,
            "Analyst": analyst,
            "PriceTarget": pt
        })
    return pd.DataFrame(rows)

def md_table(df: pd.DataFrame) -> str:
    """Return a Markdown table without requiring the 'tabulate' package."""
    if df is None or df.empty:
        return ""
    d = df.copy()
    idx_name = d.index.name or "Index"
    d = d.replace([np.inf, -np.inf], np.nan).fillna("")
    d.index = d.index.astype(str)
    for c in d.columns:
        d[c] = d[c].astype(str)
    headers = [idx_name] + list(d.columns)
    header_row = "| " + " | ".join(headers) + " |"
    sep_row    = "| " + " | ".join(["---"] * len(headers)) + " |"
    body_rows = []
    for idx, row in d.iterrows():
        vals = [idx] + [row[c] for c in d.columns]
        body_rows.append("| " + " | ".join(vals) + " |")
    return "\n".join([header_row, sep_row, *body_rows])

def color_chg(val):
    if pd.isna(val): return ""
    return "color: #0b8f3a;" if val > 0 else "color: #b30000;"

def heat_bg(v, vmax=2.0):
    """Return a green/red background-color CSS based on percent change v."""
    if pd.isna(v):
        return ""
    v = float(max(-vmax, min(vmax, v)))  # clamp
    if v >= 0:
        alpha = 0.15 + 0.55 * (v / vmax)  # 0.15..0.70
        return f"background-color: rgba(46, 204, 113, {alpha});"
    else:
        alpha = 0.15 + 0.55 * (-v / vmax)
        return f"background-color: rgba(231, 76, 60, {alpha});"

# ---------- Risk light ----------
def risk_light(es, nq, ym, vix_level, vix_chg):
    """
    Very simple score:
      +1 for each of ES, NQ, YM > +0.30%
      -1 for each of ES, NQ, YM < -0.30%
      -1 if VIX > 20, -1 if VIX change > +3%
    """
    score = 0
    for f in [es, nq, ym]:
        if not pd.isna(f):
            if f > 0.30: score += 1
            if f < -0.30: score -= 1
    if not pd.isna(vix_level) and vix_level > 20: score -= 1
    if not pd.isna(vix_chg) and vix_chg > 3.0: score -= 1

    if score >= 2: return "🟢 Green (Risk-On)"
    if score <= -2: return "🔴 Red (Risk-Off)"
    return "🟡 Yellow (Mixed)"

# ---- Sector ETFs (SPDRs) ----
SECTORS = {
    "XLC": "Comm Svcs",
    "XLY": "Cons Discr",
    "XLP": "Cons Staples",
    "XLE": "Energy",
    "XLF": "Financials",
    "XLV": "Health Care",
    "XLI": "Industrials",
    "XLB": "Materials",
    "XLRE": "Real Estate",
    "XLK": "Technology",
    "XLU": "Utilities",
}

def sector_heatmap(df: pd.DataFrame, title: str):
    """
    df expected columns: ['Last','Chg%','PM_Last','PM_%Chg','PM_Vol']
    Displays two heatmaps: Daily % change & Pre-Market % change
    """
    if df.empty:
        st.info(f"No sector data available for {title}.")
        return

    view = df.copy()
    for col in ["Last","Chg%","PM_Last","PM_%Chg","PM_Vol"]:
        if col not in view.columns:
            view[col] = np.nan

    try:
        view.index = [SECTORS.get(sym, sym) for sym in view.index]
    except Exception:
        pass

    c1, c2 = st.columns(2)

    # Daily heatmap
    with c1:
        st.subheader(f"{title} — Daily %")
        tbl = view[["Chg%","Last"]].rename(columns={"Chg%":"% Chg", "Last":"Last"})
        sty = (
            tbl.style
               .format({"% Chg": "{:+.2f}%", "Last": "{:,.2f}"})
               .applymap(lambda x: heat_bg(x, vmax=2.0), subset=["% Chg"])
        )
        st.dataframe(sty, width="stretch")

    # Pre-market heatmap
    with c2:
        st.subheader(f"{title} — Pre-Market %")
        tbl = view[["PM_%Chg","PM_Last","PM_Vol"]].rename(
            columns={"PM_%Chg":"% Chg (PM)","PM_Last":"Last (PM)","PM_Vol":"Vol (PM)"}
        )
        if "Vol (PM)" in tbl.columns:
            tbl["Vol (PM)"] = pd.to_numeric(tbl["Vol (PM)"], errors="coerce")
        sty = (
            tbl.style
               .format({
                   "% Chg (PM)": "{:+.2f}%",
                   "Last (PM)": "{:,.2f}",
                   "Vol (PM)": "{:,.0f}",
               })
               .applymap(lambda x: heat_bg(x, vmax=2.0), subset=["% Chg (PM)"])
        )
        st.dataframe(sty, width="stretch")

# ----------------- Sidebar -----------------
st.sidebar.header("⚙️ Settings")
default_symbols = "^YM=F,^ES=F,^NQ=F,^RTY=F,^VIX,^GSPC,^IXIC,CL=F,GC=F,BTC-USD,ETH-USD"
symbols = st.sidebar.text_input("Symbols (comma separated)", default_symbols)

st.sidebar.markdown("---")
st.sidebar.subheader("Watchlist")
watchlist = st.sidebar.text_input("Tickers (comma separated)", "AAPL, MSFT, NVDA, TSLA, AMD")

st.sidebar.markdown("---")
st.sidebar.subheader("Up/Downgrades Input")
ud_mode = st.sidebar.radio(
    "How will you provide U/D list?",
    ["Paste text", "Upload CSV/Excel"],
    horizontal=True,
    key="ud_mode"
)

# ----------------- Catalysts / Econ -----------------
st.header("🔥 Catalysts")
catalysts = st.text_area("Top catalysts today (one per line)")

st.header("📅 Econ Calendar")
econ = st.text_area("Key economic releases / Fed speakers")

# ----------------- Global Market Summary -----------------
st.header("🌍 Global Market Summary")

ASIA = {
    "^N225": "Nikkei 225",
    "^HSI": "Hang Seng",
    "000001.SS": "Shanghai",
    "^BSESN": "Sensex",
    "^STI": "Singapore",
    "^AORD": "ASX 200",
}
EUROPE = {
    "^FTSE": "FTSE 100",
    "^GDAXI": "DAX",
    "^FCHI": "CAC 40",
    "FTSEMIB.MI": "FTSE MIB",
    "^IBEX": "IBEX 35",
    "^STOXX": "STOXX 600",
}
USA = {
    "^DJI": "Dow",
    "^GSPC": "S&P 500",
    "^IXIC": "Nasdaq",
    "^VIX": "VIX",
    "CL=F": "Crude Oil",
    "GC=F": "Gold",
}

def show_region(title, mapping):
    syms = list(mapping.keys())
    df = fetch_daily_snapshot(syms)
    if df.empty:
        st.info(f"No data for {title}.")
        return
    df = df[["Last","Chg%"]]
    df.index = [mapping.get(i, i) for i in df.index]
    st.subheader(title)
    st.dataframe(
        df.style.format({"Last": "{:,.2f}", "Chg%": "{:+.2f}%"}).applymap(color_chg, subset=["Chg%"]),
        width="stretch"
    )

c1, c2, c3 = st.columns(3)
with c1: show_region("Asia", ASIA)
with c2: show_region("Europe", EUROPE)
with c3: show_region("U.S.", USA)

# ----------------- Futures + Premarket Panel -----------------
st.header("🌙 Overnight & 🕘 Pre-Market")

# main symbols area
syms = [s.strip() for s in symbols.split(",") if s.strip()]
daily_df = fetch_daily_snapshot(syms) if syms else pd.DataFrame()
pm_df = fetch_premarket_stats(syms) if syms else pd.DataFrame()
combo = combine_snapshot_and_pm(daily_df, pm_df) if not daily_df.empty else pd.DataFrame()

# SAFE formatted display
if not combo.empty:
    fmt = {
        "Last":            lambda v: "—" if pd.isna(v) else f"{v:,.2f}",
        "Chg%":            lambda v: "—" if pd.isna(v) else f"{v:+.2f}%",
        "PM_Last":         lambda v: "—" if pd.isna(v) else f"{v:,.2f}",
        "PM_%Chg":         lambda v: "—" if pd.isna(v) else f"{v:+.2f}%",
        "Gap_vs_Close_%":  lambda v: "—" if pd.isna(v) else f"{v:+.2f}%",
        "PM_Vol":          lambda v: "—" if pd.isna(v) else f"{int(round(v)):,}",
    }
    if "PM_Vol" in combo.columns:
        combo["PM_Vol"] = pd.to_numeric(combo["PM_Vol"], errors="coerce")

    st.dataframe(
        combo.style.format(fmt).applymap(color_chg, subset=["Chg%","PM_%Chg","Gap_vs_Close_%"]),
        width="stretch"
    )

    # Futures & VIX for risk light
    es = combo.loc["^ES=F","Chg%"] if "^ES=F" in combo.index else np.nan
    nq = combo.loc["^NQ=F","Chg%"] if "^NQ=F" in combo.index else np.nan
    ym = combo.loc["^YM=F","Chg%"] if "^YM=F" in combo.index else np.nan
    vix_level = combo.loc["^VIX","Last"] if "^VIX" in combo.index else np.nan
    vix_chg = combo.loc["^VIX","Chg%"] if "^VIX" in combo.index else np.nan
    light = risk_light(es, nq, ym, vix_level, vix_chg)

    # Bias tile from Dow futs (^YM=F) or ^DJI
    symdict = {i: daily_df.loc[i].to_dict() for i in daily_df.index} if not daily_df.empty else {}
    _bias, _pct, used_sym = bias_from_symbol(symdict)
    pct_str = "—" if pd.isna(_pct) else f"{_pct:+.2f}%"

    b1, b2 = st.columns(2)
    with b1: st.metric("🧭 Bias", _bias, f"{pct_str} ({used_sym})")
    with b2: st.metric("🚦 Risk Light", light)
else:
    st.info("Enter symbols in the sidebar (comma separated).")

# ----------------- Sector ETFs Heatmap -----------------
st.header("🗺️ Sector ETFs Heatmap")
sector_syms = list(SECTORS.keys())
sec_daily = fetch_daily_snapshot(sector_syms)
sec_pm    = fetch_premarket_stats(sector_syms)
if not sec_daily.empty:
    sec_combo = combine_snapshot_and_pm(sec_daily, sec_pm)
    sec_combo = sec_combo.loc[[s for s in sector_syms if s in sec_combo.index]]
    try:
        sector_heatmap(sec_combo, "SPDR Sectors")
    except Exception as e:
        st.warning(f"Sector heatmap unavailable: {e}")
        fmt = {
            "Last": "{:,.2f}",
            "Chg%": "{:+.2f}%",
            "PM_Last": "{:,.2f}",
            "PM_%Chg": "{:+.2f}%",
            "Gap_vs_Close_%": "{:+.2f}%",
            "PM_Vol": "{:,.0f}",
        }
        tbl = sec_combo.copy()
        if "PM_Vol" in tbl.columns:
            tbl["PM_Vol"] = pd.to_numeric(tbl["PM_Vol"], errors="coerce")
        st.dataframe(
            tbl.style.format(fmt).applymap(color_chg, subset=["Chg%","PM_%Chg","Gap_vs_Close_%"]),
            width="stretch"
        )
else:
    st.info("No sector data available right now.")

# ----------------- Watchlist Quick-Glance -----------------
st.header("👀 Watchlist Quick-Glance (Pre-Market)")
wl = [t.strip().upper() for t in watchlist.split(",") if t.strip()]
if wl:
    wl_daily = fetch_daily_snapshot(wl)
    wl_pm = fetch_premarket_stats(wl)
    wl_combo = combine_snapshot_and_pm(wl_daily, wl_pm)
    if not wl_combo.empty:
        st.dataframe(
            wl_combo.style.format({
                "Last": "{:,.2f}", "Chg%": "{:+.2f}%",
                "PM_Last": "{:,.2f}", "PM_%Chg": "{:+.2f}%",
                "Gap_vs_Close_%": "{:+.2f}%", "PM_Vol": "{:,.0f}"
            }).applymap(color_chg, subset=["Chg%","PM_%Chg","Gap_vs_Close_%"]),
            width="stretch"
        )
    else:
        st.info("No watchlist data available.")
else:
    st.caption("Add tickers to your watchlist in the sidebar.")

# ----------------- Up/Downgrades -----------------
st.header("📈 Upgrades / 📉 Downgrades")
ud_df = pd.DataFrame()
if st.session_state.get("ud_mode", ud_mode) == "Paste text":
    ud_text = st.text_area(
        "Paste upgrades/downgrades (freeform)",
        height=160,
        placeholder="AAPL Upgrade Neutral->Buy Firm: JPM Analyst: Doe PT: 210\nMSFT Downgrade Overweight->Equal-Weight Firm: MS PT: 400"
    )
    if ud_text.strip():
        ud_df = parse_updown_text(ud_text)
else:
    up = st.file_uploader(
        "Upload CSV or Excel with columns: Ticker, Action, OldRating, NewRating, Firm, Analyst, PriceTarget",
        type=["csv", "xlsx"]
    )
    if up is not None:
        try:
            if up.name.lower().endswith(".csv"):
                ud_df = pd.read_csv(up)
            else:
                ud_df = pd.read_excel(up)
        except Exception as e:
            st.error(f"Failed to read file: {e}")

if not ud_df.empty:
    expected_cols = ["Ticker","Action","OldRating","NewRating","Firm","Analyst","PriceTarget"]
    for c in expected_cols:
        if c not in ud_df.columns:
            ud_df[c] = "" if c != "PriceTarget" else np.nan
    ud_df = ud_df[expected_cols]
    st.dataframe(ud_df, width="stretch")
else:
    st.caption("No U/D entries yet.")

# ----------------- Options Ideas (quick scan) -----------------
st.header("📝 Options Ideas (Quick Chain Peek)")
c1, c2 = st.columns([1, 2])
with c1:
    opt_ticker = st.text_input("Options Ticker (e.g., AAPL)")
with c2:
    opt_sort = st.selectbox("Sort by", ["volume", "openInterest"], index=0)

if opt_ticker:
    try:
        tk = yf.Ticker(opt_ticker.strip().upper())
        exps = tk.options
        if not exps:
            st.warning("No option expirations found.")
        else:
            expiry_sel = st.selectbox("Pick expiry", exps, index=0, key="opt_expiry")
            chain = tk.option_chain(expiry_sel)
            calls = chain.calls.loc[:, ["strike", "lastPrice", "volume", "openInterest"]].sort_values(opt_sort, ascending=False).head(12)
            puts  = chain.puts.loc[:,  ["strike", "lastPrice", "volume", "openInterest"]].sort_values(opt_sort, ascending=False).head(12)

            ofmt = {
                "lastPrice": lambda v: "—" if pd.isna(v) else f"{v:,.2f}",
                "volume":    lambda v: "—" if pd.isna(v) else f"{int(round(v)):,}",
                "openInterest": lambda v: "—" if pd.isna(v) else f"{int(round(v)):,}",
                "strike":    lambda v: "—" if pd.isna(v) else f"{v:,.2f}",
            }

            col1, col2 = st.columns(2)
            with col1:
                st.subheader(f"Calls (top by {opt_sort})")
                st.dataframe(calls.style.format(ofmt), width="stretch")
            with col2:
                st.subheader(f"Puts (top by {opt_sort})")
                st.dataframe(puts.style.format(ofmt), width="stretch")
    except Exception as e:
        st.error(f"Options fetch failed: {e}")

# ----------------- Final Summary (Markdown / optional DOCX) -----------------
st.header("📝 Summary")

def safe_slice(df, cols):
    try:
        return df[cols]
    except Exception:
        return df

snap_md = ""
if not combo.empty:
    snap_md = md_table(safe_slice(combo, ["Last","Chg%","PM_Last","PM_%Chg","Gap_vs_Close_%","PM_Vol"]))

ud_md = ""
if not ud_df.empty:
    ud_md = md_table(ud_df.set_index("Ticker"))

# compute bias line safely
bias_line = "—"
try:
    if not combo.empty:
        if "^YM=F" in combo.index:
            val = float(combo.loc["^YM=F", "Chg%"])
        elif "^DJI" in combo.index:
            val = float(combo.loc["^DJI", "Chg%"])
        else:
            val = np.nan
        if not pd.isna(val):
            bias_line = f"{'Risk-On' if val > 0 else 'Risk-Off'} ({val:+.2f}%)"
except Exception:
    pass

summary = f"""
### 🔥 Catalysts
{catalysts or '-'}

### 📅 Econ
{econ or '-'}

### 🌍 Global Snapshot
- Asia, Europe, U.S. shown above

### 🌙 Overnight + 🕘 Pre-Market
{snap_md or '-'}

### 📈 Up/Downgrades
{ud_md or '-'}

### 🧭 Bias / 🚦 Risk Light
- Bias: {bias_line}
- Risk Light: {locals().get('light','—')}
""".strip()

st.markdown(summary)

# Markdown download
st.download_button(
    "⬇️ Download Report (Markdown)",
    data=summary,
    file_name=f"DMR_{datetime.now(tz=NY).date()}.md",
    mime="text/markdown"
)

# Optional: DOCX download (if python-docx available)
try:
    from docx import Document
    from docx.shared import Pt

    if st.button("⬇️ Download Report (Word .docx)"):
        doc = Document()
        doc.add_heading("Daily Market Report", level=1)
        for block in summary.split("\n\n"):
            p = doc.add_paragraph(block)
            for run in p.runs:
                run.font.size = Pt(11)
        path = f"DMR_{datetime.now(tz=NY).date()}.docx"
        doc.save(path)
        with open(path, "rb") as f:
            st.download_button(
                "Download .docx",
                f,
                file_name=path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
except Exception:
    # quietly skip Word export if package not installed
    pass

with st.expander("⚖️ Disclaimer"):
    st.write("""
The material on this report represents an assessment of the market and economic environment at a specific point in time and is not intended to be a forecast of future events, or a guarantee of future results. Forward-looking statements are subject to certain risks and uncertainties. Actual results, performance, or achievements may differ materially from those expressed or implied. Information is based on data gathered from sources believed to be reliable, but accuracy is not guaranteed and it is not intended to be used as a primary basis for investment decisions. Past performance does not guarantee future results.
""")
